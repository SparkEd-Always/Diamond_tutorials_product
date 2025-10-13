#!/usr/bin/env python3
"""Extract text from admin_tasks.docx"""

from docx import Document

def extract_docx_content(file_path):
    """Extract all text from a DOCX file"""
    doc = Document(file_path)

    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    return '\n'.join(full_text)

if __name__ == "__main__":
    import sys
    sys.stdout.reconfigure(encoding='utf-8')
    content = extract_docx_content(r"d:\Projects\sparked\study\admin_tasks.docx")
    print(content)
